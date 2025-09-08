"""
Complete document processor with async text extraction and thumbnail generation
"""
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument
from PIL import Image, ImageDraw, ImageFont
import io
import re
import time
import logging
from pathlib import Path
from uuid import UUID
from typing import Dict, Any, Optional, List, Tuple
from fastapi import UploadFile, HTTPException
from datetime import datetime

from ..services.s3_service import s3_service
from ..config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Comprehensive document processor with text extraction, metadata analysis,
    and thumbnail generation capabilities
    """
    
    def __init__(self):
        self.s3_service = s3_service
        self.supported_formats = {
            'application/pdf': self._extract_pdf_content,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_content,
            'application/msword': self._extract_docx_content,
            'text/plain': self._extract_text_content
        }

    async def process_upload(
        self, 
        file: UploadFile, 
        tenant_id: UUID,
        document_type_id: UUID,
        category_id: Optional[UUID] = None,
        tags: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Process uploaded document: validate, upload to S3, extract basic metadata
        Full text extraction happens asynchronously
        
        Args:
            file: Uploaded file
            tenant_id: Tenant identifier
            document_type_id: Document type identifier
            category_id: Optional category identifier
            tags: Optional list of tags
            
        Returns:
            Document processing results
        """
        start_time = time.time()
        
        try:
            # Validate file before processing
            self._validate_file(file)
            
            # Generate unique document ID
            from uuid import uuid4
            document_id = uuid4()
            
            # Upload to S3
            upload_result = await self.s3_service.upload_document(
                file, document_id, tenant_id, "document"
            )
            
            # Extract basic metadata quickly (without full text extraction)
            basic_metadata = await self._extract_basic_metadata(file)
            
            processing_time = time.time() - start_time
            
            result = {
                "document_id": document_id,
                "s3_key": upload_result["s3_key"],
                "original_filename": file.filename,
                "file_size": upload_result["file_size"],
                "mime_type": upload_result["content_type"],
                "file_hash": upload_result["file_hash"],
                "tenant_id": tenant_id,
                "document_type_id": document_type_id,
                "category_id": category_id,
                "tags": tags or [],
                "status": "uploaded",
                "extraction_status": "pending",
                "processing_time": processing_time,
                **basic_metadata
            }
            
            logger.info(f"Document {document_id} processed in {processing_time:.2f}s")
            return result
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Document processing failed: {str(e)}"
            )

    async def extract_text_async(
        self, 
        document_id: UUID, 
        s3_key: str, 
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Asynchronously extract text content and generate thumbnail
        This is called as a background task after upload
        
        Args:
            document_id: Document identifier
            s3_key: S3 key for the document
            mime_type: MIME type of the document
            
        Returns:
            Extraction results with text content and metadata
        """
        start_time = time.time()
        
        try:
            logger.info(f"Starting async text extraction for document {document_id}")
            
            # Get document content from S3
            file_content = await self.s3_service.get_document_content(s3_key)
            
            # Extract content based on file type
            if mime_type in self.supported_formats:
                extraction_func = self.supported_formats[mime_type]
                content_result = await extraction_func(file_content)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
            
            # Generate thumbnail
            thumbnail_result = await self._generate_thumbnail(
                file_content, mime_type, document_id
            )
            
            # Calculate text statistics
            text_stats = self._calculate_text_statistics(content_result["text"])
            
            processing_time = time.time() - start_time
            
            result = {
                "extraction_status": "completed",
                "raw_content": content_result["text"],
                "page_count": content_result.get("page_count", 1),
                "character_count": text_stats["character_count"],
                "word_count": text_stats["word_count"],
                "thumbnail_s3_key": thumbnail_result.get("s3_key"),
                "extraction_completed_at": datetime.utcnow(),
                "processing_time": processing_time,
                "metadata": {
                    **content_result.get("metadata", {}),
                    **text_stats,
                    "extraction_method": content_result.get("method", "unknown")
                }
            }
            
            logger.info(f"Text extraction completed for {document_id} in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            logger.error(f"Text extraction failed for {document_id}: {e}")
            return {
                "extraction_status": "failed",
                "extraction_error": str(e),
                "extraction_completed_at": datetime.utcnow(),
                "processing_time": time.time() - start_time
            }

    async def _extract_basic_metadata(self, file: UploadFile) -> Dict[str, Any]:
        """
        Extract basic metadata without full text processing
        
        Args:
            file: Uploaded file
            
        Returns:
            Basic metadata dictionary
        """
        try:
            # Reset file pointer
            await file.seek(0)
            content = await file.read()
            
            metadata = {
                "file_size_human": self._format_file_size(len(content)),
                "estimated_pages": 1,  # Default for non-PDF files
            }
            
            # For PDFs, quickly get page count
            if file.content_type == 'application/pdf':
                try:
                    doc = fitz.open(stream=content, filetype="pdf")
                    metadata["estimated_pages"] = doc.page_count
                    doc.close()
                except Exception as e:
                    logger.warning(f"Could not extract PDF page count: {e}")
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Basic metadata extraction failed: {e}")
            return {"file_size_human": "Unknown", "estimated_pages": 1}

    async def _extract_pdf_content(self, pdf_content: bytes) -> Dict[str, Any]:
        """
        Extract text and metadata from PDF using PyMuPDF
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Extraction result with text and metadata
        """
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
            text_parts = []
            metadata = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", "")
            }
            
            # Extract text from each page
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---")
                    text_parts.append(page_text)
            
            full_text = "\n".join(text_parts)
            page_count = doc.page_count
            
            doc.close()
            
            if not full_text.strip():
                logger.warning("No text content found in PDF")
                full_text = "[No extractable text content found]"
            
            return {
                "text": full_text,
                "page_count": page_count,
                "method": "PyMuPDF",
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise ValueError(f"PDF processing failed: {str(e)}")

    async def _extract_docx_content(self, docx_content: bytes) -> Dict[str, Any]:
        """
        Extract text from DOCX using python-docx
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Extraction result with text and metadata
        """
        try:
            doc = DocxDocument(io.BytesIO(docx_content))
            
            # Extract core properties
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or ""
            }
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if table_text:
                all_text.append("\n--- Tables ---")
                all_text.extend(table_text)
            
            full_text = "\n".join(all_text)
            
            if not full_text.strip():
                logger.warning("No text content found in DOCX")
                full_text = "[No extractable text content found]"
            
            return {
                "text": full_text,
                "page_count": 1,  # DOCX doesn't have clear page boundaries
                "method": "python-docx",
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise ValueError(f"DOCX processing failed: {str(e)}")

    async def _extract_text_content(self, text_content: bytes) -> Dict[str, Any]:
        """
        Process plain text files
        
        Args:
            text_content: Text file content as bytes
            
        Returns:
            Extraction result with text and metadata
        """
        try:
            # Try to decode as UTF-8, fallback to other encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = text_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, use UTF-8 with error handling
                text = text_content.decode('utf-8', errors='replace')
                logger.warning("Used UTF-8 with error replacement for text file")
            
            # Estimate page count based on content length
            estimated_pages = max(1, len(text) // 3000)  # ~3000 chars per page
            
            return {
                "text": text,
                "page_count": estimated_pages,
                "method": "plain_text",
                "metadata": {
                    "encoding": "utf-8",
                    "estimated_pages": estimated_pages
                }
            }
            
        except Exception as e:
            logger.error(f"Text file processing failed: {e}")
            raise ValueError(f"Text file processing failed: {str(e)}")

    async def extract_text_for_ai(self, file_content: bytes, content_type: str) -> str:
        """
        Extract text content from document for AI processing
        Reuses existing extraction methods to avoid code duplication
        
        Args:
            file_content: Document file content as bytes
            content_type: MIME type of the document
            
        Returns:
            Extracted text content as string
        """
        try:
            if content_type == 'application/pdf':
                # Reuse existing PDF extraction method
                result = await self._extract_pdf_content(file_content)
                return result.get('text', '')
                
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                # Reuse existing DOCX extraction method
                result = await self._extract_docx_content(file_content)
                return result.get('text', '')
                
            elif content_type == 'text/plain':
                # Reuse existing text extraction method
                result = await self._extract_text_content(file_content)
                return result.get('text', '')
                
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported content type: {content_type}")
                
        except Exception as e:
            logger.error(f"Text extraction for AI failed: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to extract text for AI: {str(e)}")

    async def _generate_thumbnail(
        self, 
        file_content: bytes, 
        mime_type: str, 
        document_id: UUID
    ) -> Dict[str, Any]:
        """
        Generate thumbnail for document
        
        Args:
            file_content: File content as bytes
            mime_type: MIME type of the file
            document_id: Document identifier
            
        Returns:
            Thumbnail generation result
        """
        try:
            if mime_type == 'application/pdf':
                return await self._generate_pdf_thumbnail(file_content, document_id)
            else:
                return await self._generate_generic_thumbnail(mime_type, document_id)
                
        except Exception as e:
            logger.warning(f"Thumbnail generation failed for {document_id}: {e}")
            return {"error": str(e)}

    async def _generate_pdf_thumbnail(
        self, 
        pdf_content: bytes, 
        document_id: UUID
    ) -> Dict[str, Any]:
        """
        Generate full-page preview from PDF first page
        
        Args:
            pdf_content: PDF content as bytes
            document_id: Document identifier
            
        Returns:
            Preview generation result
        """
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
            
            if doc.page_count == 0:
                raise ValueError("PDF has no pages")
            
            # Get first page
            page = doc[0]
            
            # Render page as image with very high resolution for crisp preview
            # Use higher DPI for better quality
            mat = fitz.Matrix(4, 4)  # 4x zoom for very high quality preview
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            # Convert to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # For full-page preview, we want high resolution but reasonable file size
            # Max width: 1200px for better quality, maintain aspect ratio
            max_width = 1200
            if img.width > max_width:
                ratio = max_width / img.width
                new_height = int(img.height * ratio)
                img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
            
            # Save as very high-quality JPEG
            preview_buffer = io.BytesIO()
            img.save(preview_buffer, format='JPEG', quality=95, optimize=True)
            preview_data = preview_buffer.getvalue()
            
            # Upload preview to S3
            tenant_id = UUID(str(document_id).split('-')[0] + '-0000-0000-0000-000000000000')  # Simplified
            preview_s3_key = await self.s3_service.upload_thumbnail(
                preview_data, document_id, tenant_id, "jpg"
            )
            
            doc.close()
            
            return {
                "s3_key": preview_s3_key,
                "size": len(preview_data),
                "dimensions": img.size,
                "format": "JPEG",
                "type": "full_page_preview"
            }
            
        except Exception as e:
            logger.error(f"PDF preview generation failed: {e}")
            raise e

    async def _generate_generic_thumbnail(
        self, 
        mime_type: str, 
        document_id: UUID
    ) -> Dict[str, Any]:
        """
        Generate generic thumbnail for non-PDF files
        
        Args:
            mime_type: MIME type of the file
            document_id: Document identifier
            
        Returns:
            Thumbnail generation result
        """
        try:
            # Create a simple thumbnail with file type icon
            img = Image.new('RGB', (300, 400), color='#f8fafc')
            draw = ImageDraw.Draw(img)
            
            # Determine file type icon/text
            file_type_map = {
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'DOCX',
                'application/msword': 'DOC',
                'text/plain': 'TXT'
            }
            
            file_type = file_type_map.get(mime_type, 'DOC')
            
            # Draw background
            draw.rectangle([20, 20, 280, 380], fill='white', outline='#e2e8f0', width=2)
            
            # Try to load a font (fall back to default if not available)
            try:
                font_large = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 48)
                font_small = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
            except:
                font_large = ImageFont.load_default()
                font_small = ImageFont.load_default()
            
            # Draw file type
            text_bbox = draw.textbbox((0, 0), file_type, font=font_large)
            text_width = text_bbox[2] - text_bbox[0]
            text_x = (300 - text_width) // 2
            draw.text((text_x, 180), file_type, fill='#374151', font=font_large)
            
            # Draw document icon (simple representation)
            draw.rectangle([125, 100, 175, 160], fill='#3b82f6', outline='#1e40af', width=2)
            draw.rectangle([130, 105, 170, 115], fill='white')
            draw.rectangle([130, 120, 170, 125], fill='white')
            draw.rectangle([130, 130, 170, 135], fill='white')
            
            # Save thumbnail
            thumbnail_buffer = io.BytesIO()
            img.save(thumbnail_buffer, format='JPEG', quality=85, optimize=True)
            thumbnail_data = thumbnail_buffer.getvalue()
            
            # Upload thumbnail to S3
            tenant_id = UUID(str(document_id).split('-')[0] + '-0000-0000-0000-000000000000')  # Simplified
            thumbnail_s3_key = await self.s3_service.upload_thumbnail(
                thumbnail_data, document_id, tenant_id, "jpg"
            )
            
            return {
                "s3_key": thumbnail_s3_key,
                "size": len(thumbnail_data),
                "dimensions": img.size,
                "format": "JPEG",
                "type": "generic"
            }
            
        except Exception as e:
            logger.error(f"Generic thumbnail generation failed: {e}")
            raise e

    def _calculate_text_statistics(self, text: str) -> Dict[str, int]:
        """
        Calculate basic text statistics
        
        Args:
            text: Text content
            
        Returns:
            Statistics dictionary
        """
        if not text:
            return {"character_count": 0, "word_count": 0, "line_count": 0}
        
        # Count characters (excluding whitespace)
        character_count = len(text)
        
        # Count words (split by whitespace and filter empty strings)
        words = [word for word in re.split(r'\s+', text) if word]
        word_count = len(words)
        
        # Count lines
        line_count = len(text.split('\n'))
        
        return {
            "character_count": character_count,
            "word_count": word_count,
            "line_count": line_count
        }

    def _validate_file(self, file: UploadFile):
        """
        Validate uploaded file type and size
        
        Args:
            file: Uploaded file
            
        Raises:
            HTTPException: If validation fails
        """
        # Check file type
        if file.content_type not in settings.allowed_file_types:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file.content_type}. "
                       f"Allowed types: {', '.join(settings.allowed_file_types)}"
            )
        
        # Check file size
        if file.size and file.size > settings.max_file_size:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {self._format_file_size(settings.max_file_size)}"
            )
        
        # Check filename
        if not file.filename:
            raise HTTPException(
                status_code=400,
                detail="Filename is required"
            )
        
        # Check for potentially dangerous file extensions
        dangerous_extensions = ['.exe', '.bat', '.cmd', '.scr', '.pif', '.com']
        file_ext = Path(file.filename).suffix.lower()
        if file_ext in dangerous_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed for security reasons: {file_ext}"
            )

    def _format_file_size(self, size_bytes: int) -> str:
        """
        Format file size in human-readable format
        
        Args:
            size_bytes: Size in bytes
            
        Returns:
            Formatted size string
        """
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        while size_bytes >= 1024 and i < len(size_names) - 1:
            size_bytes /= 1024
            i += 1
        
        return f"{size_bytes:.1f} {size_names[i]}"

    async def detect_document_type(self, text: str, filename: str, mime_type: str) -> str:
        """
        Auto-detect document type based on content, filename, and MIME type
        
        Args:
            text: Document text content
            filename: Original filename
            mime_type: File MIME type
            
        Returns:
            Detected document type name
        """
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Define document type keywords and patterns
        type_patterns = {
            'invoice': {
                'keywords': ['invoice', 'bill', 'payment', 'amount due', 'total', 'tax', 'subtotal', 'invoice number', 'billing'],
                'filename_patterns': ['invoice', 'bill', 'receipt'],
                'score_multiplier': 1.0
            },
            'contract': {
                'keywords': ['agreement', 'contract', 'terms', 'conditions', 'party', 'whereas', 'hereby', 'executed', 'binding'],
                'filename_patterns': ['contract', 'agreement', 'terms'],
                'score_multiplier': 1.0
            },
            'insurance_policy': {
                'keywords': ['policy', 'coverage', 'premium', 'deductible', 'claim', 'insured', 'beneficiary', 'policy number'],
                'filename_patterns': ['policy', 'insurance', 'coverage'],
                'score_multiplier': 1.0
            }
        }
        
        # Score each document type
        type_scores = {}
        for doc_type, patterns in type_patterns.items():
            score = 0
            
            # Score based on content keywords
            for keyword in patterns['keywords']:
                if keyword in text_lower:
                    score += text_lower.count(keyword) * patterns['score_multiplier']
            
            # Score based on filename patterns (higher weight)
            for pattern in patterns['filename_patterns']:
                if pattern in filename_lower:
                    score += 3 * patterns['score_multiplier']
            
            type_scores[doc_type] = score
        
        # Return type with highest score, or default based on file type
        if type_scores and max(type_scores.values()) > 0:
            return max(type_scores, key=type_scores.get)
        
        # Default fallback based on MIME type
        if mime_type == 'application/pdf':
            return 'contract'  # PDFs often contracts
        elif 'word' in mime_type or 'document' in mime_type:
            return 'contract'  # Word docs often contracts
        else:
            return 'invoice'  # Default fallback

    async def detect_document_category(self, text: str, filename: str) -> Optional[str]:
        """
        Simple heuristic-based document category detection
        
        Args:
            text: Document text content
            filename: Original filename
            
        Returns:
            Suggested category name or None
        """
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Define category keywords
        category_keywords = {
            'Invoice': ['invoice', 'bill', 'payment', 'amount due', 'total', 'tax', 'subtotal'],
            'Contract': ['agreement', 'contract', 'terms', 'conditions', 'party', 'whereas', 'hereby'],
            'Insurance': ['policy', 'coverage', 'premium', 'deductible', 'claim', 'insured', 'beneficiary'],
            'Legal': ['court', 'legal', 'lawsuit', 'attorney', 'law', 'judgment', 'plaintiff', 'defendant'],
            'Personal': ['birth certificate', 'passport', 'driver', 'license', 'personal', 'identity']
        }
        
        # Score each category
        category_scores = {}
        for category, keywords in category_keywords.items():
            score = 0
            for keyword in keywords:
                if keyword in text_lower:
                    score += text_lower.count(keyword)
                if keyword in filename_lower:
                    score += 2  # Filename matches are weighted higher
            category_scores[category] = score
        
        # Return category with highest score (if above threshold)
        if category_scores:
            best_category = max(category_scores, key=category_scores.get)
            if category_scores[best_category] > 0:
                return best_category
        
        return None


# Global document processor instance
document_processor = DocumentProcessor()
